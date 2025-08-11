
from io import BytesIO
from typing import Optional, Dict, Any, List
import re, json, os

# local imports
try:
    from docx import Document
except Exception:
    Document = None

from pathlib import Path
from gemini_client import generate_completion

DATA_REF = Path(__file__).parent / "data" / "adgm_references.txt"

def _ensure_docx_available():
    if Document is None:
        raise RuntimeError("python-docx is required but not installed. Please pip install python-docx")

def extract_text_from_docx_bytes(doc_bytes: bytes) -> str:
    _ensure_docx_available()
    bio = BytesIO(doc_bytes)
    doc = Document(bio)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text for cell in row.cells])
            parts.append(row_text)
    return "\n".join(parts)

DOC_TYPE_KEYWORDS = {
    "incorporation": ["resolution", "incorporation", "memorandum", "articles of association", "board resolution"],
    "employment_contract": ["employment", "employee", "employer", "salary", "probation", "termination"],
    "agreement": ["agreement", "party", "witnesseth", "whereas"]
}

def classify_doc_type(full_text: str) -> str:
    lower = full_text.lower()
    scores = {k:0 for k in DOC_TYPE_KEYWORDS}
    for k, kws in DOC_TYPE_KEYWORDS.items():
        for kw in kws:
            if kw in lower:
                scores[k] += 1
    best = max(scores, key=lambda k: scores[k])
    if scores[best] == 0:
        return "unknown"
    return best

INCORP_CHECKLIST = [
    "memorandum of association", "articles of association", "board resolution", "shareholder agreement", "registered address"
]

def checklist_missing_sections(full_text: str, doc_type: str) -> List[str]:
    lower = full_text.lower()
    missing = []
    if doc_type == "incorporation":
        for item in INCORP_CHECKLIST:
            if item not in lower:
                missing.append(item)
    return missing

RED_FLAG_KEYWORDS = [
    "power of attorney", "undisclosed beneficiary", "subject to approval", "without prejudice", "as necessary",
    "to be decided", "to be determined", "sole discretion", "undisclosed beneficiary"
]

def detect_red_flags(full_text: str) -> List[str]:
    lower = full_text.lower()
    found = []
    for kw in RED_FLAG_KEYWORDS:
        if kw in lower:
            found.append(kw)
    if re.search(r"\b0\s*-\s*100\b", full_text):
        found.append("suspicious percentage range 0-100")
    return found

def insert_comments_into_docx_bytes(doc_bytes: bytes, notes: Dict[str, str]) -> bytes:
    _ensure_docx_available()
    bio = BytesIO(doc_bytes)
    doc = Document(bio)
    doc.add_page_break()
    doc.add_paragraph("=== Automated Review Notes ===")
    for key, note in notes.items():
        doc.add_paragraph(f"{key}: {note}")
    out = BytesIO()
    doc.save(out)
    return out.getvalue()

# SIMPLE retrieval: score snippets by keyword overlap
def retrieve_adgm_citations(text: str, top_k: int = 3) -> List[Dict[str,str]]:
    if not DATA_REF.exists():
        return []
    corpus = DATA_REF.read_text(encoding="utf-8")
    paras = [p.strip() for p in corpus.split("\n\n") if p.strip()][:500]
    text_lower = text.lower()
    scored = []
    for p in paras:
        score = 0
        lower_p = p.lower()
        for word in set(text_lower.split()):
            if len(word) > 4 and word in lower_p:
                score += 1
        if score > 0:
            scored.append((score, p))
    scored.sort(key=lambda x: x[0], reverse=True)
    results = [{"score": s, "snippet": p[:800]} for s,p in scored[:top_k]]
    return results

def process_docx(doc_bytes: Optional[bytes], query: Optional[str]) -> Dict[str, Any]:
    if doc_bytes is None and (not query):
        return {"error": "No document or query provided."}
    full_text = ""
    if doc_bytes:
        try:
            full_text = extract_text_from_docx_bytes(doc_bytes)
        except Exception as e:
            return {"error": f"Failed to parse .docx: {e}"}
    else:
        full_text = query or ""

    doc_type = classify_doc_type(full_text)
    missing = checklist_missing_sections(full_text, doc_type)
    red_flags = detect_red_flags(full_text)

    notes = {}
    if missing:
        notes["Missing sections"] = ", ".join(missing)
    if red_flags:
        notes["Red flags"] = ", ".join(red_flags)
    if not notes:
        notes["Notes"] = "No immediate issues detected by heuristic checks."

    marked_bytes = None
    if doc_bytes:
        try:
            marked_bytes = insert_comments_into_docx_bytes(doc_bytes, notes)
        except Exception as e:
            marked_bytes = None

    # Retrieve ADGM citations (simple)
    citations = []
    try:
        citations = retrieve_adgm_citations(full_text or query or "")
    except Exception:
        citations = []

    # Create a Gemini prompt to get context-aware suggestions
    gemini_result = None
    try:
        prompt = f"""You are an assistant that reviews corporate legal documents for ADGM compliance.
Given the document text below, provide (1) a short summary of main issues, and (2) three practical suggestions to fix them. Output as JSON with keys: summary, suggestions.
Document text:
""" + (full_text or query or "")[:4000]
        gemini_result = generate_completion(prompt, context="\n\n".join([c['snippet'] for c in citations]))
    except Exception as e:
        gemini_result = {"text": f"Gemini not available: {e}"}

    summary = {
        "doc_type": doc_type,
        "missing_sections": missing,
        "red_flags": red_flags,
        "notes": notes
    }

    return {
        "doc_type": doc_type,
        "missing_sections": missing,
        "red_flags": red_flags,
        "notes": notes,
        "marked_docx_bytes": marked_bytes,
        "summary": summary,
        "top_citations": citations,
        "gemini": gemini_result
    }

if __name__ == "__main__":
    print("pipeline module updated with Gemini integration (mock-capable).")


def generate_report_blob(summary: dict) -> bytes:
    """
    Generates a simple text report from the summary and returns it as bytes.
    You can replace this with DOCX or PDF generation later if needed.
    """
    report_lines = []
    report_lines.append(f"Document Type: {summary.get('doc_type', 'Unknown')}")
    report_lines.append(f"Missing Sections: {', '.join(summary.get('missing_sections', [])) or 'None'}")
    report_lines.append(f"Red Flags: {', '.join(summary.get('red_flags', [])) or 'None'}")
    report_lines.append("Notes:")
    for k, v in summary.get('notes', {}).items():
        report_lines.append(f"  {k}: {v}")
    return "\n".join(report_lines).encode("utf-8")
